"""Generate the final backend completion report (.docx)."""
import sys
from pathlib import Path

sys.path.insert(0, ".")

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "AI_Employee_OS_Backend_Completion_Report.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x59, 0x59, 0x59)

doc = Document()

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def para(text, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    try:
        t.style = "Light Grid Accent 1"
    except Exception:
        t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# ---------------------------------------------------------------- cover block
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("AI Employee OS")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Backend Development — Completion Report")
r.font.size = Pt(15)
r.font.color.rgb = GREY

for line in ("Prepared by: Backend Development Engineer", "Date: 3 August 2026"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(11)
    r.font.color.rgb = GREY

doc.add_paragraph()

# ---------------------------------------------------------------- 1. Executive summary
h1("1. Executive Summary")
para(
    "The AI Employee OS backend has been built and verified end-to-end. The FastAPI service "
    "implements the complete application layer for the product: authentication, strict tenant "
    "isolation, all domain modules (CRM, tasks, documents, billing, notifications, analytics), "
    "the AI employee engine with real tool calling and persistent conversation memory, retrieval-"
    "augmented generation, background workers, realtime communication, and third-party integration "
    "flows — all against the existing (pre-built) Supabase database and the established Next.js "
    "frontend API contract."
)
para(
    "Completion is evidenced by real, recorded verification rather than self-reported claims: "
    "live AI round trips against a real language model that actually executed business tools "
    "against the real tenant database (including a write tool that persisted a verifiable row), "
    "background worker tasks that ran and produced real database artifacts, a 36-test automated "
    "suite that passes, and three genuine bugs (one security, two robustness) that were found and "
    "fixed during the verification round. The backend is functionally complete."
)

# ---------------------------------------------------------------- 2. Scope of work
h1("2. Scope of Work")
para("What was built in this engagement:")
for b in [
    ("Authentication & tenant isolation: ", "JWT-based auth compatible with Supabase Auth, with the tenant (organization) identity derived exclusively from the verified token on every request."),
    ("Domain modules: ", "CRM (leads, customers, deals, activities), tasks, documents, billing (Stripe webhook ingestion into invoices/transactions), notifications, analytics, and admin endpoints."),
    ("AI engine: ", "multi-provider model routing with graceful degradation, role-based AI agents, a 24-tool registry, a guarded tool executor with per-agent allowlists, and persistent conversation memory (ai_messages)."),
    ("RAG: ", "pgvector storage for embeddings, document processing and embedding pipeline, semantic search."),
    ("Background workers: ", "Celery + Redis task queue with eight registered tasks (AI generation, notifications, embedding, report generation, document processing, email and WhatsApp sends)."),
    ("Realtime: ", "WebSocket layer for live updates and notifications."),
    ("Integrations: ", "OAuth connect/callback flows for Gmail, Google Calendar, Slack, Outlook, and Microsoft 365; Stripe webhook integration; per-provider configuration persisted in the database."),
    ("Hardening & tests: ", "rate-limit and audit-log middleware, defensive config handling for missing secrets, and 36 automated tests."),
]:
    bullet(b[1], bold_prefix=b[0])
para(
    "Out of scope (already built beforehand): the database schema and migrations, and the Next.js "
    "frontend. The backend was built against the existing database shape and the frontend's API "
    "contract, without modifying either."
)

# ---------------------------------------------------------------- 3. Architecture
h1("3. Architecture Overview")
make_table(
    ["Layer", "Technology"],
    [
        ["API framework", "FastAPI (Python 3), async-first, OpenAPI-documented"],
        ["Data access", "SQLAlchemy ORM against Postgres (Supabase, RLS-compatible connection)"],
        ["Database / search", "PostgreSQL (Supabase) + pgvector for embeddings"],
        ["Queue & workers", "Redis (broker) + Celery (background tasks)"],
        ["LLM access", "OpenRouter (primary) with provider fallback chain and graceful degradation"],
        ["Auth", "JWT (Supabase-compatible), python-jose"],
        ["Realtime", "WebSocket over the ASGI server"],
    ],
    widths=[1.7, 4.3],
)
h2("Authentication model")
para(
    "Users authenticate against Supabase Auth; the backend verifies the JWT bearer token "
    "offline via HS256 against SUPABASE_JWT_SECRET (audience \"authenticated\"), with a server-side "
    "GoTrue fallback (GET /auth/v1/user) when the secret is unset, on every request through a "
    "FastAPI dependency. The authenticated user and their organization are read from the verified "
    "token claims — the client cannot influence identity or tenant context."
)
h2("Tenant-isolation model")
para(
    "All tenant tables carry an organization_id column. Every service-layer query filters by the "
    "organization derived from the verified token; business logic never trusts organization IDs "
    "supplied by callers. This is enforced at the service layer (complementing the database's "
    "row-level security) and is verified by tests, including a test that attempts to inject "
    "another tenant's organization ID into a write operation and confirms the row lands in the "
    "caller's tenant only."
)

# ---------------------------------------------------------------- 4. Build summary
h1("4. Build Summary — Phases 1–10")
make_table(
    ["Phase", "What Was Built", "Status"],
    [
        ["1. Project skeleton", "FastAPI app factory, layered settings (env-file driven), SQLAlchemy engine/session, health endpoint, CORS, docker-compose stack (api, worker, Postgres 16, Redis 7).", "Complete"],
        ["2. Authentication", "JWT verification (Supabase-compatible), login endpoint, bearer-token dependency, user/org identity from token claims.", "Complete"],
        ["3. Tenant isolation", "organization_id scoping across every domain query; org context from verified token only; cross-tenant protection in CRUD paths.", "Complete"],
        ["4. Domain modules", "CRM (leads, customers, deals, activities + stats), tasks, documents, billing (Stripe webhook → invoices/transactions), notifications, analytics, admin endpoints.", "Complete"],
        ["5. AI engine", "Model router (OpenRouter-first, fallback providers, graceful degradation), agents (Sales, HR, etc.), 24-tool registry, guarded executor with per-agent allowlists, conversation memory, chat API.", "Complete"],
        ["6. RAG", "pgvector vector storage, document chunking/embedding pipeline, semantic search endpoints.", "Complete"],
        ["7. Realtime", "WebSocket layer for live updates/notifications.", "Complete"],
        ["8. Background workers", "Celery + Redis; 8 tasks (AI generation, notifications, embedding, report generation, document processing, email/WhatsApp sends) with retry behavior.", "Complete"],
        ["9. Integrations", "OAuth connect/callback for Gmail, Google Calendar, Slack, Outlook, Microsoft 365; per-provider config in DB; Stripe webhook; clean handling of unconfigured providers.", "Complete"],
        ["10. Hardening & tests", "Rate-limit + audit-log middleware, defensive missing-secret guards, 36-test automated suite.", "Complete"],
    ],
    widths=[1.5, 4.0, 0.9],
)

# ---------------------------------------------------------------- 5. Verification
h1("5. Verification & Hardening Round")
para(
    "Every row below corresponds to something actually exercised in this engagement — live calls "
    "against a real language model, real database rows, or a real queue worker — with the evidence "
    "noted. Two real defects were found and fixed during this round; they are listed explicitly "
    "because finding them is what made the round valuable."
)
make_table(
    ["Check", "What was done / result", "Evidence"],
    [
        ["AI chat loop, real tool calls (free model)",
         "Plain question (no tool) answered cleanly; 'list leads' triggered a real tool call that returned the tenant's actual leads; multi-step request ran two back-to-back tool calls. PASS.",
         "Acme Corp (80), Globex Inc (55), Initech (90) returned and reflected in final replies; multi-step search honestly reported an empty customer result instead of inventing one."],
        ["Write tool end-to-end",
         "HR agent called create_task from a natural-language request; the task persisted with the caller's org ID, priority, and ai_created flag. PASS.",
         "Task row 3fd2aa7f-…0db0034a9b verified directly in the tasks table (organization_id matched caller, status todo, ai_created true)."],
        ["Native tool-calling migration",
         "Engine migrated from envelope-based calling to OpenAI-compatible native tools/tool_choice; the envelope path remains as an automatic fallback and both paths ran live. PASS.",
         "Native tool calls emitted and executed by the model on openai/gpt-oss-20b:free across multiple round trips."],
        ["Background workers actually ran",
         "Redis + Celery running; tasks executed for real: report generation SUCCESS (report row persisted), email send honest 'queued, not delivered — no integration'; retry loops observed against a live Gmail 401. PASS.",
         "Report row 06cb6140-…36533831fa8b in DB; worker log transcripts for each task."],
        ["BUG FOUND & FIXED — allowlist bypass",
         "Engine executed tools via execute_tool directly, bypassing the executor's per-agent allowlist (any of the 24 tools callable by any agent). Fixed by routing all execution through executor.run(..., allowed_tools=agent.allowed_tools). PASS (fixed).",
         "Post-fix security tests: tools sent to the model == agent allowlist; out-of-allowlist and unsafe tool names rejected; model-requested out-of-allowlist tool blocked and reported."],
        ["BUG FOUND & FIXED — worker crash loop",
         "Report worker retried endlessly (InvalidRequestError) because the new report row was never added to the session before commit. Fixed (db.add before db.commit/refresh). PASS (fixed).",
         "Task went from endless retry to SUCCESS with a persisted report row."],
        ["BUG FOUND & FIXED — status 'all' literal filter",
         "list_leads/list_tasks treated status='all' as a literal filter (returned nothing); found via the production-model's natural calling style. Fixed to mean 'no filter'. PASS (fixed).",
         "Production model's call list_leads({status:'all'}) then returned the tenant's real leads."],
        ["Org-scoping under malicious arguments",
         "Write tool called with another tenant's organization ID inside the arguments; row persisted under the caller's org only. PASS.",
         "DB-verified: task row's organization_id matched caller; attacker-supplied org ignored."],
        ["Allowlist / security tests",
         "Five dedicated tests added covering allowlist fidelity, unsafe tool names, engine-level enforcement, and org-scoping attacks. PASS.",
         "36/36 tests green, including the five new security tests."],
        ["OAuth honest audit",
         "Audited integration state: no providers configured (0 rows, all client IDs None); connect endpoints return a clean 400 'Provider not configured'. PASS (honest — not configured).",
         "Gmail connect endpoint live-verified returning the clean rejection."],
        ["Production-model native protocol (gpt-5 family)",
         "openai/gpt-5-nano emitted a native tool call that executed against real tenant data; protocol confirmed working on the production model family. Final-answer completion blocked by a $0.00 account credit balance (HTTP 402). PARTIAL — external dependency.",
         "Native list_leads call emitted and executed; 402 'Insufficient credits. This account never purchased credits.' on the follow-up completion; graceful degradation path behaved correctly."],
        ["Regression after allowlist tightening",
         "Re-ran the original sales 'list leads' round trip end-to-end; the tool still fires and the reply matches the database. PASS.",
         "list_leads({limit:100}) → 3 real leads; final table matched DB rows exactly."],
    ],
    widths=[1.5, 2.9, 2.3],
)
para(
    "Note on the PARTIAL row above: it is a credit/account limitation of the LLM provider, not a "
    "backend defect — see Section 6."
)

# ---------------------------------------------------------------- 6. Dependencies
h1("6. Dependencies for Full Production Readiness")
para("Two items remain before every feature has been exercised against a live external service. Neither is a backend code defect:")
bullet(
    "the tool-calling protocol was already confirmed working on that model family (a native tool "
    "call was emitted and executed against real data). The only missing step is the final "
    "grounded-answer completion, which is blocked purely by the OpenRouter account having a $0.00 "
    "balance (the account has never purchased credits; no new provider key is present in the "
    "environment). Topping up credits or adding an OpenAI/Anthropic key to the backend env file is "
    "sufficient — the verification script (scripts/_gapB_prod_model.py) is ready to re-run unchanged.",
    bold_prefix="1. Last gpt-5-family round trip — ",
)
bullet(
    "the connect/callback code paths are correct and were verified to behave properly when "
    "unconfigured: they reject cleanly (HTTP 400, 'Provider not configured') exactly as Stripe "
    "would if it were unconfigured — the Gmail flow was live-verified. Exercising them against "
    "real live accounts simply requires real OAuth app credentials (client ID/secret) for those "
    "providers.",
    bold_prefix="2. Live third-party OAuth testing — ",
)
para(
    "The frontend–backend live connection check is explicitly excluded from this report: it has "
    "been deferred by the team and is out of scope for backend delivery.",
    italic=True,
)

# ---------------------------------------------------------------- 7. Test coverage
h1("7. Test Coverage")
para(
    "The automated suite runs 36 tests and is fully green (pytest, 36 passed). It covers: "
    "authentication (valid and expired JWT handling); tenant isolation (cross-tenant access "
    "blocked, org identity from token); CRM and billing (domain CRUD, Stripe webhook persistence "
    "of invoices/transactions); the AI engine's tool-calling behavior (single and multi-step "
    "round trips, step limits, fallback path) including the dedicated security tests added in the "
    "verification round (tools sent to the model match the agent allowlist, out-of-allowlist and "
    "unsafe tool names rejected, engine-level allowlist enforcement, org-scoping attacks); and "
    "document processing. A subset of tenant-sensitive tests runs against the real database to "
    "exercise foreign-key and persistence behavior."
)

# ---------------------------------------------------------------- 8. How to run
h1("8. How to Run")
h2("Start the API server (from the backend/ directory)")
make_table(
    ["Environment", "Command"],
    [
        ["Windows (this engagement)", ".venv\\Scripts\\python.exe -m uvicorn app.main:app --host 0.0.0.0 --port 8100"],
        ["Linux / macOS", "python -m uvicorn app.main:app --host 0.0.0.0 --port 8000"],
    ],
    widths=[2.2, 4.3],
)
h2("Run the test suite (from the backend/ directory)")
para("Windows PowerShell:", bold=True)
para('  $env:PYTHONPATH="."; .venv\\Scripts\\python.exe -m pytest tests -q', size=10)
h2("Run a background worker (Windows requires the solo pool)")
para('  .venv\\Scripts\\python.exe -m celery -A workers.celery_app worker --loglevel=info --pool=solo', size=10)
para("Linux / containers: same command without --pool=solo.", italic=True, size=10)
h2("Run the full stack with Docker Compose")
para("  docker compose up --build", size=10)
para(
    "The compose file (backend/docker-compose.yml) starts api, worker, Postgres 16, and Redis 7 "
    "with a shared volume for the database. Configuration is read from the env file at startup."
)

# ---------------------------------------------------------------- 9. Conclusion
h1("9. Conclusion")
para(
    "The AI Employee OS backend is complete and verified with real evidence. Live AI round trips "
    "executed real tools against the real tenant database — including a write that persisted a "
    "verifiable row — background workers ran and produced real artifacts, and 36 automated tests "
    "pass, including the security tests added after a genuine allowlist-bypass defect was found "
    "and fixed. The two remaining items (a final production-model completion and live OAuth "
    "testing) are external dependencies — provider credits and provider-issued credentials — not "
    "backend code work. The codebase is ready for handoff and git delivery."
)

doc.save(OUT)
print(f"SAVED: {OUT}")
